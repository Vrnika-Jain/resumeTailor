import streamlit as st
from groq import Groq
import os
import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fpdf import FPDF

st.set_page_config(
    page_title="Resume Tailor",
    page_icon="📄",
    layout="centered"
)

st.markdown("""
    <style>
    .block-container { max-width: 720px; padding-top: 2rem; }
    .stTextArea textarea { font-size: 14px; }
    .stDownloadButton button { width: 100%; }
    </style>
""", unsafe_allow_html=True)

BLUE     = (27, 79, 138)
DARK     = (20, 20, 20)
GRAY     = (100, 100, 100)
BLUE_HEX = "1B4F8A"


def sanitize(text: str) -> str:
    return (text
        .replace("\u2013", "-").replace("\u2014", "-")
        .replace("\u2018", "'").replace("\u2019", "'")
        .replace("\u201c", '"').replace("\u201d", '"')
        .replace("\u2022", "-").replace("\u2026", "...")
        .replace("\u00e9", "e").replace("\u00e0", "a")
        .replace("\u00e1", "a").replace("\u00e8", "e")
        .replace("\u00ea", "e").replace("\u00eb", "e")
        .replace("\u00ef", "i").replace("\u00ee", "i")
        .replace("\u00f4", "o").replace("\u00f6", "o")
        .replace("\u00fa", "u").replace("\u00fb", "u")
        .replace("\u00fc", "u").replace("\u00f1", "n")
        .replace("\u00e7", "c").replace("\u00b7", "-")
        .replace("\u00ae", "(R)").replace("\u00a9", "(C)")
        .replace("\u2122", "(TM)")
    )


def tailor_with_groq(resume: str, jd: str, api_key: str) -> str:
    key = (
        api_key.strip()
        or st.secrets.get("GROQ_API_KEY", "")
        or os.environ.get("GROQ_API_KEY", "")
    )
    if not key:
        raise ValueError("No API key provided.")

    client = Groq(api_key=key)

    prompt = f"""You are an expert resume writer and ATS optimization specialist.

TASK: Rewrite the candidate's resume to perfectly target the job description below.
Keep ALL sections — do not remove any. Tailor and reword content to match the job.

STRICT FORMATTING RULES (follow exactly, no exceptions):
- Line 1: Candidate full name only
- Line 2: email | phone | linkedin_url | location | website_url (only include what exists in the resume)
- Line 3: Gender: value | Date of birth: value | Nationality: value
- Then sections in this exact order:
  WORK EXPERIENCE
  INTERNSHIP EXPERIENCE
  EDUCATION AND TRAINING
  RELEVANT PROJECTS
  SKILL SET
  LANGUAGES
  CERTIFICATIONS
  ACCOMPLISHMENTS
  PUBLICATIONS

- Under WORK EXPERIENCE and INTERNSHIP EXPERIENCE:
  Each entry:
  DATE: start - end
  ROLE: Job Title | Organization | City, Country
  - bullet point
  - bullet point

- Under EDUCATION AND TRAINING:
  Each entry:
  DATE: start - end
  ROLE: Degree/Certificate | Institution | City, Country
  GRADE: grade value

- Under RELEVANT PROJECTS (include only projects relevant to the job description, check technologies used):
  Each entry:
  PROJECT: Project Name | date range
  DESC: one line description tailored to job
  TECH: technologies
  LINK: url (only if present in resume)

- Under SKILL SET:
  CAT: Category Name
  ITEMS: skill1, skill2, skill3

- Under LANGUAGES:
  LANG: Language | proficiency level

- Under CERTIFICATIONS:
  CERT: certification name | issuer

- Under ACCOMPLISHMENTS:
  ACCOMP: description

- Under PUBLICATIONS (only if present in resume):
  PUB: title | url

INSTRUCTIONS:
1. Extract the most important keywords and skills from the job description.
2. Weave these naturally into bullets and skills.
3. Rewrite bullets to emphasize accomplishments relevant to this role.
4. Only include projects whose TECH stack is relevant to the job description.
5. Do NOT invent experience - only reframe existing experience.
6. Use strong action verbs. Be concise and impactful.
7. If the candidate resume contains a website, portfolio link, or publication, always include them.
8. Keep Internship Experience section until experience exceeds 3 years.
9. Keep Accomplishments section until experience exceeds 3 years.

CANDIDATE RESUME:
{resume}

JOB DESCRIPTION:
{jd}

Output ONLY the tailored resume. No commentary, no explanations, no markdown.
"""

    msg = client.chat.completions.create(
        model="meta-llama/llama-4-scout-17b-16e-instruct",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=3000,
        temperature=0.7
    )
    return msg.choices[0].message.content.strip()


def generate_cover_letter(resume: str, jd: str, api_key: str) -> str:
    key = (
        api_key.strip()
        or st.secrets.get("GROQ_API_KEY", "")
        or os.environ.get("GROQ_API_KEY", "")
    )
    if not key:
        raise ValueError("No API key provided.")

    client = Groq(api_key=key)

    prompt = f"""You are an expert cover letter writer.

TASK: Write a formal and professional cover letter tailored to the job description below.
Base it entirely on the candidate's resume. Do not invent experience.

FORMATTING RULES (follow exactly):
- Line 1: Candidate full name
- Line 2: email | phone | location
- Line 3: Today's date written out e.g. 10 April 2026
- Line 4: blank
- Line 5: Hiring Manager title and company name extracted from job description
- Line 6: blank
- Line 7: Dear Hiring Manager,
- Line 8: blank
- Then body paragraphs separated by blank lines
- End with:
  Yours sincerely,

  [Candidate full name]

INSTRUCTIONS:
1. Opening paragraph: Express genuine interest in the specific role and company.
   Mention the exact job title. Keep it concise and compelling.
2. Middle paragraphs: Highlight 2-3 most relevant experiences or achievements
   from the resume that directly match the job requirements.
   Use specific numbers or outcomes where possible.
   Reference keywords from the job description naturally.
3. Closing paragraph: Summarise why you are a strong fit.
   Express enthusiasm for an interview. Keep it confident but not arrogant.
4. Tone: Formal, professional, and human. No buzzwords or clichés.
5. Length: Let the content decide — write as many paragraphs as needed
   to make a compelling case, but keep it to one page maximum.
6. Do NOT use bullet points anywhere in the cover letter.
7. Do NOT repeat the resume — tell a story instead.

CANDIDATE RESUME:
{resume}

JOB DESCRIPTION:
{jd}

Output ONLY the cover letter. No commentary, no explanations, no markdown.
"""

    msg = client.chat.completions.create(
        model="meta-llama/llama-4-scout-17b-16e-instruct",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=1500,
        temperature=0.7
    )
    return msg.choices[0].message.content.strip()


def parse_resume(text: str) -> dict:
    lines = [l.rstrip() for l in text.split("\n")]
    data = {
        "name": "",
        "contact": "",
        "meta": "",
        "sections": []
    }
    if lines:
        data["name"] = lines[0].strip()
    if len(lines) > 1:
        data["contact"] = lines[1].strip()
    if len(lines) > 2:
        data["meta"] = lines[2].strip()

    current_section = None
    headers = {
        "WORK EXPERIENCE", "INTERNSHIP EXPERIENCE", "EDUCATION AND TRAINING",
        "RELEVANT PROJECTS", "SKILL SET", "LANGUAGES",
        "CERTIFICATIONS", "ACCOMPLISHMENTS", "PUBLICATIONS"
    }
    for line in lines[3:]:
        s = line.strip()
        if s.upper() in headers:
            current_section = {"title": s.upper(), "lines": []}
            data["sections"].append(current_section)
        elif current_section and s:
            current_section["lines"].append(s)
    return data


def build_pdf(data: dict) -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_margins(20, 20, 20)
    pdf.set_auto_page_break(auto=True, margin=20)
    W = pdf.w - 40

    # Name
    pdf.set_xy(20, 20)
    pdf.set_font("Helvetica", "B", 20)
    pdf.set_text_color(*DARK)
    pdf.cell(W, 10, sanitize(data["name"]), ln=True)

    # Meta: Gender, Date of birth, Nationality
    meta = data.get("meta", "")
    if meta:
        meta = (meta
            .replace("GENDER:", "Gender:")
            .replace("DOB:", "Date of birth:")
            .replace("NATIONALITY:", "Nationality:")
        )
        parts = [p.strip() for p in meta.split("|")]
        meta_line = "    ".join(parts)
        pdf.set_x(20)
        pdf.set_font("Helvetica", "B", 8.5)
        pdf.set_text_color(*DARK)
        pdf.multi_cell(W, 5, sanitize(meta_line))

    # Contact — wrap to next line if too long
    contact = data.get("contact", "")
    contact_parts = [c.strip() for c in contact.split("|")]
    pdf.set_font("Helvetica", "", 8.5)
    pdf.set_text_color(*BLUE)

    char_width  = 2.1
    line_parts  = []
    current_w   = 0

    for part in contact_parts:
        part_w = len(part) * char_width + 15
        if current_w + part_w > W and line_parts:
            pdf.set_x(20)
            pdf.cell(W, 5, "  |  ".join(line_parts), ln=True)
            line_parts  = [part]
            current_w   = part_w
        else:
            line_parts.append(part)
            current_w += part_w

    if line_parts:
        pdf.set_x(20)
        pdf.cell(W, 5, "  |  ".join(line_parts), ln=True)

    # Header underline
    pdf.ln(2)
    pdf.set_draw_color(*BLUE)
    pdf.set_line_width(0.8)
    pdf.line(20, pdf.get_y(), pdf.w - 20, pdf.get_y())
    pdf.ln(5)

    # Sections
    for sec in data["sections"]:
        pdf.set_x(20)
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(*DARK)
        pdf.cell(W, 6, sanitize(sec["title"]), ln=True)
        pdf.set_draw_color(*BLUE)
        pdf.set_line_width(0.4)
        pdf.line(20, pdf.get_y(), pdf.w - 20, pdf.get_y())
        pdf.ln(3)

        i = 0
        lines = sec["lines"]
        while i < len(lines):
            line = lines[i]
            pdf.set_x(20)

            if line.startswith("DATE:"):
                date_val = sanitize(line.replace("DATE:", "").strip())
                role_val = ""
                if i + 1 < len(lines) and lines[i+1].startswith("ROLE:"):
                    role_val = sanitize(lines[i+1].replace("ROLE:", "").strip())
                    i += 2
                else:
                    i += 1
                pdf.set_font("Helvetica", "I", 8.5)
                pdf.set_text_color(*GRAY)
                pdf.cell(42, 5, date_val, ln=False)
                pdf.set_font("Helvetica", "B", 10)
                pdf.set_text_color(*DARK)
                pdf.multi_cell(W - 42, 5, role_val)

            elif line.startswith("ROLE:"):
                role_val = sanitize(line.replace("ROLE:", "").strip())
                pdf.set_font("Helvetica", "B", 10)
                pdf.set_text_color(*DARK)
                pdf.set_x(62)
                pdf.multi_cell(W - 42, 5, role_val)
                i += 1

            elif line.startswith("GRADE:"):
                grade_val = sanitize(line.replace("GRADE:", "").strip())
                pdf.set_font("Helvetica", "", 9)
                pdf.set_text_color(*DARK)
                pdf.set_x(62)
                pdf.multi_cell(W - 42, 5, "Grade: " + grade_val)
                i += 1

            elif line.startswith("- "):
                pdf.set_font("Helvetica", "", 9.5)
                pdf.set_text_color(*DARK)
                pdf.set_x(62)
                pdf.multi_cell(W - 42, 5, "-  " + sanitize(line[2:]))
                i += 1

            elif line.startswith("PROJECT:"):
                pdf.set_font("Helvetica", "B", 10)
                pdf.set_text_color(*DARK)
                pdf.set_x(20)
                pdf.multi_cell(W, 5, sanitize(line.replace("PROJECT:", "").strip()))
                i += 1

            elif line.startswith("DESC:"):
                pdf.set_font("Helvetica", "", 9.5)
                pdf.set_text_color(*DARK)
                pdf.set_x(20)
                pdf.multi_cell(W, 5, sanitize(line.replace("DESC:", "").strip()))
                i += 1

            elif line.startswith("TECH:"):
                pdf.set_font("Helvetica", "I", 9)
                pdf.set_text_color(*GRAY)
                pdf.set_x(20)
                pdf.multi_cell(W, 5, "Technologies: " + sanitize(line.replace("TECH:", "").strip()))
                i += 1

            elif line.startswith("LINK:"):
                pdf.set_font("Helvetica", "I", 9)
                pdf.set_text_color(*BLUE)
                pdf.set_x(20)
                pdf.multi_cell(W, 5, sanitize(line.replace("LINK:", "").strip()))
                i += 1

            elif line.startswith("CAT:"):
                cat_val = sanitize(line.replace("CAT:", "").strip())
                items_val = ""
                if i + 1 < len(lines) and lines[i+1].startswith("ITEMS:"):
                    items_val = sanitize(lines[i+1].replace("ITEMS:", "").strip())
                    i += 2
                else:
                    i += 1
                pdf.set_font("Helvetica", "B", 9.5)
                pdf.set_text_color(*DARK)
                pdf.set_x(20)
                pdf.cell(42, 5, cat_val, ln=False)
                pdf.set_font("Helvetica", "", 9.5)
                pdf.set_text_color(*DARK)
                pdf.multi_cell(W - 42, 5, items_val)

            elif line.startswith("LANG:"):
                lang_val = sanitize(line.replace("LANG:", "").strip())
                parts = lang_val.split("|")
                pdf.set_font("Helvetica", "B", 9.5)
                pdf.set_text_color(*DARK)
                pdf.set_x(20)
                pdf.cell(42, 5, parts[0].strip() if parts else "", ln=False)
                pdf.set_font("Helvetica", "", 9.5)
                pdf.multi_cell(W - 42, 5, parts[1].strip() if len(parts) > 1 else "")
                i += 1
                continue

            elif line.startswith("CERT:"):
                pdf.set_font("Helvetica", "", 9.5)
                pdf.set_text_color(*DARK)
                pdf.set_x(20)
                pdf.multi_cell(W, 5, "-  " + sanitize(line.replace("CERT:", "").strip()))
                i += 1

            elif line.startswith("ACCOMP:"):
                pdf.set_font("Helvetica", "", 9.5)
                pdf.set_text_color(*DARK)
                pdf.set_x(20)
                pdf.multi_cell(W, 5, "-  " + sanitize(line.replace("ACCOMP:", "").strip()))
                i += 1

            elif line.startswith("PUB:"):
                pub_val = sanitize(line.replace("PUB:", "").strip())
                parts = pub_val.split("|")
                title = parts[0].strip() if parts else pub_val
                link  = parts[1].strip() if len(parts) > 1 else ""
                pdf.set_font("Helvetica", "", 9.5)
                pdf.set_text_color(*DARK)
                pdf.set_x(20)
                pdf.multi_cell(W, 5, "-  " + title)
                if link:
                    pdf.set_font("Helvetica", "I", 9)
                    pdf.set_text_color(*BLUE)
                    pdf.set_x(20)
                    pdf.multi_cell(W, 5, "   " + link)
                i += 1

            else:
                pdf.set_font("Helvetica", "", 9.5)
                pdf.set_text_color(*DARK)
                pdf.multi_cell(W, 5, sanitize(line))
                i += 1

            pdf.set_x(20)

        pdf.ln(4)

    return bytes(pdf.output())


def remove_cell_borders(tbl):
    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border in ['top','left','bottom','right','insideH','insideV']:
                b = OxmlElement(f'w:{border}')
                b.set(qn('w:val'), 'none')
                tcBorders.append(b)
            tcPr.append(tcBorders)


def add_section_header(doc, title):
    head_p = doc.add_paragraph()
    head_p.paragraph_format.space_before = Pt(8)
    head_p.paragraph_format.space_after  = Pt(2)
    pPr = head_p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '8')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), BLUE_HEX)
    pBdr.append(bot)
    pPr.append(pBdr)
    hr = head_p.add_run(title)
    hr.bold = True
    hr.font.size = Pt(11)
    hr.font.color.rgb = RGBColor(*DARK)


def build_docx(data: dict) -> bytes:
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin   = Cm(2.0)
        sec.right_margin  = Cm(2.0)

    # Name
    name_p = doc.add_paragraph()
    name_p.paragraph_format.space_after = Pt(2)
    nr = name_p.add_run(data["name"])
    nr.bold = True
    nr.font.size = Pt(20)
    nr.font.color.rgb = RGBColor(*DARK)

    # Meta
    meta = data.get("meta", "")
    if meta:
        meta = (meta
            .replace("GENDER:", "Gender:")
            .replace("DOB:", "Date of birth:")
            .replace("NATIONALITY:", "Nationality:")
        )
        meta_p = doc.add_paragraph()
        meta_p.paragraph_format.space_after = Pt(2)
        parts = [p.strip() for p in meta.split("|")]
        for idx, part in enumerate(parts):
            if ":" in part:
                label, val = part.split(":", 1)
                mr = meta_p.add_run(label.strip() + ": ")
                mr.bold = True
                mr.font.size = Pt(9)
                mr.font.color.rgb = RGBColor(*DARK)
                vr = meta_p.add_run(val.strip())
                vr.font.size = Pt(9)
                vr.font.color.rgb = RGBColor(*DARK)
            else:
                mr = meta_p.add_run(part)
                mr.font.size = Pt(9)
                mr.font.color.rgb = RGBColor(*DARK)
            if idx < len(parts) - 1:
                sr = meta_p.add_run("    ")
                sr.font.size = Pt(9)

    # Contact
    contact_p = doc.add_paragraph()
    contact_p.paragraph_format.space_after = Pt(4)
    contact_parts = [c.strip() for c in data["contact"].split("|")]
    for idx, part in enumerate(contact_parts):
        cr = contact_p.add_run(part.strip())
        cr.font.size = Pt(9)
        cr.font.color.rgb = RGBColor(*BLUE)
        if idx < len(contact_parts) - 1:
            sep = contact_p.add_run("   |   ")
            sep.font.size = Pt(9)
            sep.font.color.rgb = RGBColor(*GRAY)

    # Header underline
    div_p = doc.add_paragraph()
    div_p.paragraph_format.space_after = Pt(6)
    pPr = div_p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '16')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), BLUE_HEX)
    pBdr.append(bot)
    pPr.append(pBdr)

    # Sections
    for sec in data["sections"]:
        add_section_header(doc, sec["title"])

        i = 0
        lines = sec["lines"]
        while i < len(lines):
            line = lines[i]

            if line.startswith("DATE:"):
                date_val = line.replace("DATE:", "").strip()
                role_val = ""
                if i + 1 < len(lines) and lines[i+1].startswith("ROLE:"):
                    role_val = lines[i+1].replace("ROLE:", "").strip()
                    i += 2
                else:
                    i += 1
                tbl = doc.add_table(rows=1, cols=2)
                tbl.columns[0].width = Cm(3.8)
                tbl.columns[1].width = Cm(13.2)
                remove_cell_borders(tbl)
                dp = tbl.rows[0].cells[0].paragraphs[0]
                dr = dp.add_run(date_val)
                dr.italic = True
                dr.font.size = Pt(8.5)
                dr.font.color.rgb = RGBColor(*GRAY)
                rp = tbl.rows[0].cells[1].paragraphs[0]
                rr = rp.add_run(role_val)
                rr.bold = True
                rr.font.size = Pt(10)
                rr.font.color.rgb = RGBColor(*DARK)

            elif line.startswith("GRADE:"):
                grade_val = line.replace("GRADE:", "").strip()
                tbl = doc.add_table(rows=1, cols=2)
                tbl.columns[0].width = Cm(3.8)
                tbl.columns[1].width = Cm(13.2)
                remove_cell_borders(tbl)
                tbl.rows[0].cells[0].paragraphs[0].add_run("")
                gp = tbl.rows[0].cells[1].paragraphs[0]
                gr = gp.add_run("Grade: " + grade_val)
                gr.font.size = Pt(9)
                gr.font.color.rgb = RGBColor(*DARK)
                i += 1

            elif line.startswith("ROLE:"):
                role_val = line.replace("ROLE:", "").strip()
                tbl = doc.add_table(rows=1, cols=2)
                tbl.columns[0].width = Cm(3.8)
                tbl.columns[1].width = Cm(13.2)
                remove_cell_borders(tbl)
                tbl.rows[0].cells[0].paragraphs[0].add_run("")
                rp = tbl.rows[0].cells[1].paragraphs[0]
                rr = rp.add_run(role_val)
                rr.bold = True
                rr.font.size = Pt(10)
                rr.font.color.rgb = RGBColor(*DARK)
                i += 1

            elif line.startswith("- "):
                tbl = doc.add_table(rows=1, cols=2)
                tbl.columns[0].width = Cm(3.8)
                tbl.columns[1].width = Cm(13.2)
                remove_cell_borders(tbl)
                tbl.rows[0].cells[0].paragraphs[0].add_run("")
                bp = tbl.rows[0].cells[1].paragraphs[0]
                bp.paragraph_format.space_after = Pt(1)
                br = bp.add_run("-  " + line[2:])
                br.font.size = Pt(9.5)
                br.font.color.rgb = RGBColor(*DARK)
                i += 1

            elif line.startswith("PROJECT:"):
                pp = doc.add_paragraph()
                pp.paragraph_format.space_before = Pt(4)
                pr = pp.add_run(line.replace("PROJECT:", "").strip())
                pr.bold = True
                pr.font.size = Pt(10)
                pr.font.color.rgb = RGBColor(*DARK)
                i += 1

            elif line.startswith("DESC:"):
                dp = doc.add_paragraph()
                dr = dp.add_run(line.replace("DESC:", "").strip())
                dr.font.size = Pt(9.5)
                dr.font.color.rgb = RGBColor(*DARK)
                i += 1

            elif line.startswith("TECH:"):
                tp = doc.add_paragraph()
                tr = tp.add_run("Technologies: " + line.replace("TECH:", "").strip())
                tr.italic = True
                tr.font.size = Pt(9)
                tr.font.color.rgb = RGBColor(*GRAY)
                i += 1

            elif line.startswith("LINK:"):
                lp = doc.add_paragraph()
                lr = lp.add_run(line.replace("LINK:", "").strip())
                lr.italic = True
                lr.font.size = Pt(9)
                lr.font.color.rgb = RGBColor(*BLUE)
                i += 1

            elif line.startswith("CAT:"):
                cat_val = line.replace("CAT:", "").strip()
                items_val = ""
                if i + 1 < len(lines) and lines[i+1].startswith("ITEMS:"):
                    items_val = lines[i+1].replace("ITEMS:", "").strip()
                    i += 2
                else:
                    i += 1
                tbl = doc.add_table(rows=1, cols=2)
                tbl.columns[0].width = Cm(3.8)
                tbl.columns[1].width = Cm(13.2)
                remove_cell_borders(tbl)
                cp = tbl.rows[0].cells[0].paragraphs[0]
                cr = cp.add_run(cat_val)
                cr.bold = True
                cr.font.size = Pt(9.5)
                cr.font.color.rgb = RGBColor(*DARK)
                ip = tbl.rows[0].cells[1].paragraphs[0]
                ir = ip.add_run(items_val)
                ir.font.size = Pt(9.5)
                ir.font.color.rgb = RGBColor(*DARK)

            elif line.startswith("LANG:"):
                lang_val = line.replace("LANG:", "").strip()
                parts = lang_val.split("|")
                tbl = doc.add_table(rows=1, cols=2)
                tbl.columns[0].width = Cm(3.8)
                tbl.columns[1].width = Cm(13.2)
                remove_cell_borders(tbl)
                lp = tbl.rows[0].cells[0].paragraphs[0]
                lr = lp.add_run(parts[0].strip() if parts else "")
                lr.bold = True
                lr.font.size = Pt(9.5)
                lr.font.color.rgb = RGBColor(*DARK)
                vp = tbl.rows[0].cells[1].paragraphs[0]
                vr = vp.add_run(parts[1].strip() if len(parts) > 1 else "")
                vr.font.size = Pt(9.5)
                vr.font.color.rgb = RGBColor(*DARK)
                i += 1

            elif line.startswith("CERT:"):
                cp = doc.add_paragraph()
                cp.paragraph_format.space_after = Pt(1)
                cr = cp.add_run("-  " + line.replace("CERT:", "").strip())
                cr.font.size = Pt(9.5)
                cr.font.color.rgb = RGBColor(*DARK)
                i += 1

            elif line.startswith("ACCOMP:"):
                ap = doc.add_paragraph()
                ap.paragraph_format.space_after = Pt(1)
                ar = ap.add_run("-  " + line.replace("ACCOMP:", "").strip())
                ar.font.size = Pt(9.5)
                ar.font.color.rgb = RGBColor(*DARK)
                i += 1

            elif line.startswith("PUB:"):
                pub_val = line.replace("PUB:", "").strip()
                parts = pub_val.split("|")
                title = parts[0].strip() if parts else pub_val
                link  = parts[1].strip() if len(parts) > 1 else ""
                pp = doc.add_paragraph()
                pp.paragraph_format.space_after = Pt(1)
                pr = pp.add_run("-  " + title)
                pr.font.size = Pt(9.5)
                pr.font.color.rgb = RGBColor(*DARK)
                if link:
                    lr = pp.add_run("  " + link)
                    lr.italic = True
                    lr.font.size = Pt(9)
                    lr.font.color.rgb = RGBColor(*BLUE)
                i += 1

            else:
                p = doc.add_paragraph()
                r = p.add_run(line)
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(*DARK)
                i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_cover_letter_pdf(text: str) -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_margins(25, 25, 25)
    pdf.set_auto_page_break(auto=True, margin=25)
    W = pdf.w - 50

    lines = text.split("\n")

    for idx, line in enumerate(lines):
        line = sanitize(line.strip())

        if idx == 0:
            pdf.set_font("Helvetica", "B", 16)
            pdf.set_text_color(*DARK)
            pdf.set_x(25)
            pdf.multi_cell(W, 8, line)

        elif idx == 1:
            pdf.set_font("Helvetica", "", 9)
            pdf.set_text_color(*BLUE)
            pdf.set_x(25)
            pdf.multi_cell(W, 5, line)
            pdf.ln(2)
            pdf.set_draw_color(*BLUE)
            pdf.set_line_width(0.5)
            pdf.line(25, pdf.get_y(), pdf.w - 25, pdf.get_y())
            pdf.ln(6)

        elif idx == 2:
            pdf.set_font("Helvetica", "", 9.5)
            pdf.set_text_color(*GRAY)
            pdf.set_x(25)
            pdf.multi_cell(W, 5, line)
            pdf.ln(2)

        elif line == "":
            pdf.ln(4)

        elif line.startswith("Dear"):
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(*DARK)
            pdf.set_x(25)
            pdf.multi_cell(W, 6, line)
            pdf.ln(2)

        elif line.startswith("Yours sincerely") or line.startswith("Yours faithfully"):
            pdf.ln(4)
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(*DARK)
            pdf.set_x(25)
            pdf.multi_cell(W, 6, line)

        else:
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(*DARK)
            pdf.set_x(25)
            pdf.multi_cell(W, 6, line)

    return bytes(pdf.output())


def build_cover_letter_docx(text: str) -> bytes:
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    lines = text.split("\n")

    for idx, line in enumerate(lines):
        line = line.strip()

        if idx == 0:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(line)
            r.bold = True
            r.font.size = Pt(16)
            r.font.color.rgb = RGBColor(*DARK)

        elif idx == 1:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(line)
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(*BLUE)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bot = OxmlElement('w:bottom')
            bot.set(qn('w:val'), 'single')
            bot.set(qn('w:sz'), '8')
            bot.set(qn('w:space'), '1')
            bot.set(qn('w:color'), BLUE_HEX)
            pBdr.append(bot)
            pPr.append(pBdr)

        elif idx == 2:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(4)
            r = p.add_run(line)
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(*GRAY)

        elif line == "":
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

        elif line.startswith("Dear"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(6)
            r = p.add_run(line)
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(*DARK)

        elif line.startswith("Yours sincerely") or line.startswith("Yours faithfully"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(line)
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(*DARK)

        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(line)
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(*DARK)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


st.title("📄 Resume Tailor")
st.caption("Paste your resume and a job description — get a tailored resume and cover letter instantly.")
st.divider()

has_key = bool(
    st.secrets.get("GROQ_API_KEY", "")
    or os.environ.get("GROQ_API_KEY", "")
)

if has_key:
    api_key = ""
    st.info("API key loaded — you are good to go.", icon="✅")
else:
    api_key = st.text_input(
        "Groq API key",
        type="password",
        placeholder="gsk_...",
        help="Free API key from console.groq.com — no credit card needed."
    )

col1, col2 = st.columns(2)
with col1:
    resume_input = st.text_area(
        "Your current resume",
        height=260,
        placeholder="Paste your resume here — work experience, skills, education, everything."
    )
with col2:
    jd_input = st.text_area(
        "Job description",
        height=260,
        placeholder="Paste the full job posting here."
    )

if st.button("✨ Tailor my resume + cover letter", type="primary", use_container_width=True):
    if not has_key and not api_key.strip():
        st.error("Please enter your Groq API key.")
    elif not resume_input.strip():
        st.error("Please paste your resume.")
    elif not jd_input.strip():
        st.error("Please paste the job description.")
    else:
        with st.spinner("Tailoring your resume and writing cover letter..."):
            try:
                tailored_text = tailor_with_groq(resume_input, jd_input, api_key)
                cover_text    = generate_cover_letter(resume_input, jd_input, api_key)
                data          = parse_resume(tailored_text)
                pdf_bytes     = build_pdf(data)
                docx_bytes    = build_docx(data)
                cl_pdf_bytes  = build_cover_letter_pdf(cover_text)
                cl_docx_bytes = build_cover_letter_docx(cover_text)

                st.success("Your resume and cover letter are ready!")

                st.subheader("Resume preview")
                st.text_area("", tailored_text, height=300)

                st.subheader("Cover letter preview")
                st.text_area("", cover_text, height=300)

                st.subheader("Download resume")
                dl1, dl2 = st.columns(2)
                with dl1:
                    st.download_button(
                        "⬇ Resume — Word (.docx)",
                        data=docx_bytes,
                        file_name="tailored_resume.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                with dl2:
                    st.download_button(
                        "⬇ Resume — PDF",
                        data=pdf_bytes,
                        file_name="tailored_resume.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )

                st.subheader("Download cover letter")
                cl1, cl2 = st.columns(2)
                with cl1:
                    st.download_button(
                        "⬇ Cover Letter — Word (.docx)",
                        data=cl_docx_bytes,
                        file_name="cover_letter.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                with cl2:
                    st.download_button(
                        "⬇ Cover Letter — PDF",
                        data=cl_pdf_bytes,
                        file_name="cover_letter.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )

            except Exception as e:
                st.error(f"Something went wrong: {e}")
